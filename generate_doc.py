#!/bin/env python3
"""YouTrack Module tests"""
from youtrack import Issues
from docx import Document


##################
####  INPUTS  ####
##################

#domain
domain = input('Enter domain: ')
#username
username = input('Enter your username: ')
#token
i_token = input('Enter token path: ')
auth_token = open(i_token).read()
auth_token = auth_token.replace('\n', '')

#project ID
proj = input('Enter project ID: ')

#####################
#### DEFINITIONS ####
#####################

#site
site = 'https://' + domain + '.myjetbrains.com'
#file name for output document
fnamedoc = '%s.docx' % proj

###################
#### INSTANCES ####
###################

#output document
doc = Document()
#issues
issues = Issues(site, auth_token, project=proj)


print('Output file: ' + fnamedoc)
print('Saving issues...')


############################
####  DOCUMENT WRITING  ####
############################

#iterates over issues
for i in issues:
    #adds issue ID, summary and description into new paragraph
    doc.add_paragraph(
    'Issue ID: '    + i['id']          + '\n' +
    'Summary: '     + i['summary']     + '\n' + 
    'Description: ' + i['description'] + '\n')
    print('Issue %s (%s)' % (i['id'], i['summary']))
    #iterates over attachments
    for e in issues.attmap:
        #when attachment belongs to issue:
        if(e[0] == i['id']):
            #download attachment
            att = issues.download_attachment(e)
            #define path for saving attachment
            filename = 'images/' + e[0] + '___' + e[1]
            #open file path
            image = open(filename, 'wb')
            #saves image into file and closes file
            image.write(att)
            image.close()
            #try to add picture to document. if there is an exception, continue to the next case (can happen)
            try:
                doc.add_picture(str(filename))
            except:
                continue

###########################
####  DOCUMENT SAVING  ####
###########################

doc.save(fnamedoc)